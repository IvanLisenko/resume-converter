import asyncio
from pathlib import Path
from uuid import uuid4

from docx import Document
from fastapi.testclient import TestClient

from app.core.security import hash_password
from app.db.session import async_session_maker
from app.main import create_app
from app.models.enums import UserRole
from app.models.user import User


def _run(coro):
    return asyncio.run(coro)


async def _create_user(email: str, password: str, role: UserRole, is_active: bool = True) -> User:
    async with async_session_maker() as session:
        user = User(
            email=email,
            full_name="Интеграционный пользователь",
            password_hash=hash_password(password),
            role=role,
            is_active=is_active,
        )
        session.add(user)
        await session.commit()
        await session.refresh(user)
        return user


def _auth_headers(client: TestClient, email: str, password: str) -> dict[str, str]:
    response = client.post(
        "/api/v1/auth/login",
        json={
            "email": email,
            "password": password,
        },
    )
    assert response.status_code == 200
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


def _create_trive_docx(path: Path) -> None:
    document = Document()
    for paragraph in [
        "Иван Иванов Иванович",
        "Java Developer",
        "Опыт работы: 5 лет",
        "Навыки",
        "Java, Spring, PostgreSQL",
        "- REST API",
        "О себе",
        "Опытный Java-разработчик.",
        "Образование",
        "МГУ, Прикладная информатика (2015)",
        "Знание языков",
        "Русский — родной",
        "Английский — B2",
        "Опыт работы",
        "Senior Java Developer на проекте в сфере финтех",
        "(январь 2020 — май 2024) — 4 года 5 месяцев",
        "Описание проекта: Платформа обработки платежей",
        "Задачи: Разработка backend-сервисов",
        "Достижения: ускорил обработку платежей",
        "Команда проекта: 4 backend, 2 QA",
        "Стек: Java, Spring Framework (Spring Boot, Web), PostgreSQL",
    ]:
        document.add_paragraph(paragraph)
    document.save(path)


def _create_partner_template_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("{{ candidate.full_name }}")
    document.add_paragraph("{%p if candidate.position_line %}")
    document.add_paragraph("{{ candidate.position_line }}")
    document.add_paragraph("{%p endif %}")
    document.add_paragraph("{%p if has_summary %}")
    document.add_paragraph("{{ summary_heading }}")
    document.add_paragraph("{{ summary }}")
    document.add_paragraph("{%p endif %}")
    document.add_paragraph("{%p if has_skills %}")
    document.add_paragraph("{{ skills_heading }}")
    document.add_paragraph("{{ skills.primary_text }}")
    document.add_paragraph("{%p endif %}")
    document.add_paragraph("{%p for project in experience %}")
    document.add_paragraph("{{ project.title }}")
    document.add_paragraph("{{ project.description }}")
    document.add_paragraph("{%p for task in project.responsibilities %}")
    document.add_paragraph("{{ task }}")
    document.add_paragraph("{%p endfor %}")
    document.add_paragraph("{%p endfor %}")
    document.save(path)


def _create_unknown_variable_template_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("{{ candidate.birth_date }}")
    document.save(path)


def test_auth_login_me_and_admin_access_control():
    client = TestClient(create_app())
    suffix = uuid4().hex
    admin_email = f"admin-{suffix}@example.com"
    recruiter_email = f"recruiter-{suffix}@example.com"
    password = "Password123!"
    admin = _run(_create_user(admin_email, password, UserRole.ADMIN))
    _run(_create_user(recruiter_email, password, UserRole.RECRUITER))

    admin_headers = _auth_headers(client, admin_email, password)
    me_response = client.get("/api/v1/auth/me", headers=admin_headers)
    assert me_response.status_code == 200
    assert me_response.json()["id"] == str(admin.id)
    assert me_response.json()["role"] == UserRole.ADMIN

    bad_login = client.post(
        "/api/v1/auth/login",
        json={
            "email": admin_email,
            "password": "wrong-password",
        },
    )
    assert bad_login.status_code == 401
    assert bad_login.json()["code"] == "auth.unauthorized"

    recruiter_headers = _auth_headers(client, recruiter_email, password)
    forbidden = client.get("/api/v1/admin/users", headers=recruiter_headers)
    assert forbidden.status_code == 403
    assert forbidden.json()["code"] == "auth.forbidden"


def test_admin_can_create_employee_and_change_access():
    client = TestClient(create_app())
    suffix = uuid4().hex
    admin_email = f"admin-users-{suffix}@example.com"
    employee_email = f"employee-{suffix}@example.com"
    password = "Password123!"
    _run(_create_user(admin_email, password, UserRole.ADMIN))
    headers = _auth_headers(client, admin_email, password)

    create_response = client.post(
        "/api/v1/admin/users",
        headers=headers,
        json={
            "email": employee_email,
            "full_name": "Новый сотрудник",
            "password": "Employee123!",
            "role": "RECRUITER",
        },
    )
    assert create_response.status_code == 201
    user_id = create_response.json()["id"]
    assert create_response.json()["is_active"] is True

    role_response = client.patch(
        f"/api/v1/admin/users/{user_id}/role",
        headers=headers,
        json={"role": "ADMIN"},
    )
    assert role_response.status_code == 200
    assert role_response.json()["role"] == "ADMIN"

    block_response = client.patch(f"/api/v1/admin/users/{user_id}/block", headers=headers)
    assert block_response.status_code == 200
    assert block_response.json()["is_active"] is False

    unblock_response = client.patch(f"/api/v1/admin/users/{user_id}/unblock", headers=headers)
    assert unblock_response.status_code == 200
    assert unblock_response.json()["is_active"] is True

    list_response = client.get("/api/v1/admin/users", headers=headers)
    assert list_response.status_code == 200
    assert any(user["id"] == user_id for user in list_response.json())


def test_resume_extract_validation_errors(tmp_path):
    client = TestClient(create_app())
    suffix = uuid4().hex
    recruiter_email = f"extract-{suffix}@example.com"
    password = "Password123!"
    _run(_create_user(recruiter_email, password, UserRole.RECRUITER))
    headers = _auth_headers(client, recruiter_email, password)

    not_docx = client.post(
        "/api/v1/resumes/extract",
        headers=headers,
        files={"file": ("resume.txt", b"text", "text/plain")},
    )
    assert not_docx.status_code == 400
    assert not_docx.json()["code"] == "resume.invalid_format"

    too_large_path = tmp_path / "large.docx"
    too_large_path.write_bytes(b"x" * (10 * 1024 * 1024 + 1))
    too_large = client.post(
        "/api/v1/resumes/extract",
        headers=headers,
        files={
            "file": (
                "large.docx",
                too_large_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert too_large.status_code == 413
    assert too_large.json()["code"] == "resume.file_too_large"


def test_partner_template_upload_generation_and_operation_logs(tmp_path):
    client = TestClient(create_app())
    suffix = uuid4().hex
    admin_email = f"admin-flow-{suffix}@example.com"
    password = "Password123!"
    _run(_create_user(admin_email, password, UserRole.ADMIN))
    headers = _auth_headers(client, admin_email, password)

    partner_response = client.post(
        "/api/v1/admin/partners",
        headers=headers,
        json={
            "code": f"partner-{suffix[:12]}",
            "name": "Тестовый партнёр",
            "description": "Интеграционный тест",
        },
    )
    assert partner_response.status_code == 201
    partner_id = partner_response.json()["id"]

    public_partners = client.get("/api/v1/partners", headers=headers)
    assert public_partners.status_code == 200
    assert any(partner["id"] == partner_id for partner in public_partners.json())

    invalid_template_path = tmp_path / "invalid-template.docx"
    _create_unknown_variable_template_docx(invalid_template_path)
    invalid_template = client.post(
        f"/api/v1/admin/partners/{partner_id}/templates",
        headers=headers,
        files={
            "file": (
                "invalid-template.docx",
                invalid_template_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert invalid_template.status_code == 400
    assert invalid_template.json()["code"] == "template.unknown_variable"

    template_path = tmp_path / "partner-template.docx"
    _create_partner_template_docx(template_path)
    template_response = client.post(
        f"/api/v1/admin/partners/{partner_id}/templates",
        headers=headers,
        files={
            "file": (
                "partner-template.docx",
                template_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert template_response.status_code == 201
    assert template_response.json()["partner_id"] == partner_id
    assert template_response.json()["version"] == 1
    assert template_response.json()["is_active"] is True
    assert "candidate.full_name" in template_response.json()["variables_schema"]["variables"]

    templates_response = client.get(
        f"/api/v1/admin/partners/{partner_id}/templates",
        headers=headers,
    )
    assert templates_response.status_code == 200
    assert len(templates_response.json()) == 1

    resume_path = tmp_path / "resume.docx"
    _create_trive_docx(resume_path)
    extract_response = client.post(
        "/api/v1/resumes/extract",
        headers=headers,
        files={
            "file": (
                "resume.docx",
                resume_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert extract_response.status_code == 200
    resume = extract_response.json()["resume"]
    assert resume["candidate"]["full_name"] == "Иван Иванов Иванович"

    generate_response = client.post(
        "/api/v1/resumes/generate",
        headers=headers,
        json={
            "partnerId": partner_id,
            "resume": resume,
        },
    )
    assert generate_response.status_code == 200
    assert generate_response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    generated_path = tmp_path / "generated.docx"
    generated_path.write_bytes(generate_response.content)
    generated = Document(generated_path)
    text = "\n".join(paragraph.text for paragraph in generated.paragraphs)
    assert "Иван Иванов Иванович" in text
    assert "Java Developer" in text
    assert "Платформа обработки платежей" in text
    assert "{{" not in text
    assert "}}" not in text

    logs_response = client.get("/api/v1/admin/operation-logs?limit=20", headers=headers)
    assert logs_response.status_code == 200
    allowed_log_keys = {
        "id",
        "user_id",
        "partner_id",
        "operation_type",
        "status",
        "error_code",
        "duration_ms",
        "created_at",
    }
    assert all(set(item) == allowed_log_keys for item in logs_response.json())
    operation_types = [item["operation_type"] for item in logs_response.json()]
    assert "CREATE_PARTNER" in operation_types
    assert "UPLOAD_TEMPLATE" in operation_types
    assert "EXTRACT_RESUME" in operation_types
    assert "GENERATE_RESUME" in operation_types
