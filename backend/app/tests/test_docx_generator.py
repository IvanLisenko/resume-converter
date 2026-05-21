from docx import Document

from app.documents.docx_generator import DocxGenerationService
from app.schemas.resume import ResumeCandidate, ResumeData, ResumeSkills


def test_docx_generator_renders_template_without_unresolved_variables(tmp_path):
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "generated.docx"

    template = Document()
    template.add_paragraph("{{ candidate.full_name }}")
    template.add_paragraph("{{ candidate.position }}")
    template.add_paragraph("{{ skills.primary_text }}")
    template.save(template_path)

    resume = ResumeData(
        candidate=ResumeCandidate(
            full_name="Иван Иванов Иванович",
            position="Java Developer",
        ),
        skills=ResumeSkills(primary=["Java", "Spring"]),
    )

    DocxGenerationService().generate(template_path, resume, output_path)

    generated = Document(output_path)
    text = "\n".join(paragraph.text for paragraph in generated.paragraphs)
    assert "Иван Иванов Иванович" in text
    assert "Java Developer" in text
    assert "Java, Spring" in text
    assert "{{" not in text
    assert "}}" not in text
